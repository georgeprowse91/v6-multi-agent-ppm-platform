#!/usr/bin/env python
"""Generate prototype metadata JSON from the repo's source docs (DOCX/Markdown).

Outputs:
- apps/prototype/data/requirements.json
- apps/prototype/data/agents.json

This is useful if you update the docs and want the prototype UI to reflect the latest text.
"""

from __future__ import annotations

import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document


def repo_root() -> Path:
    # apps/prototype/scripts/generate_metadata.py -> repo root
    return Path(__file__).resolve().parents[3]


def data_dir() -> Path:
    return Path(__file__).resolve().parents[1] / "data"


def parse_prd(prd_path: Path) -> Dict[str, Any]:
    prd = Document(str(prd_path))
    functional_sections = []
    current = None
    for p in prd.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        style = p.style.name if p.style else ""
        if style == "Heading 3":
            if current:
                functional_sections.append(current)
            current = {"title": txt, "requirements": []}
        elif current and style not in ("Heading 1", "Heading 2"):
            current["requirements"].append(txt)
    if current:
        functional_sections.append(current)

    # Non-functional requirements: paragraphs under Heading 2 containing 'Non' and 'Functional'
    nf_start = None
    for i, p in enumerate(prd.paragraphs):
        if p.style.name == "Heading 2" and "Non" in p.text and "Functional" in p.text:
            nf_start = i
            break
    non_functional = []
    if nf_start is not None:
        i = nf_start + 1
        while i < len(prd.paragraphs):
            p = prd.paragraphs[i]
            if p.style.name == "Heading 2" and i != nf_start:
                break
            txt = (p.text or "").strip()
            if txt and p.style.name not in ("Heading 1", "Heading 2", "Heading 3"):
                non_functional.append(txt)
            i += 1

    return {
        "source": str(prd_path.relative_to(repo_root())),
        "functional_domains": functional_sections,
        "non_functional_requirements": non_functional,
    }


def parse_agent_doc(path: Path) -> Dict[str, Any]:
    content = path.read_text(encoding="utf-8")
    title = None
    sections = []
    current = None
    buf: List[str] = []
    for raw_line in content.splitlines():
        txt = raw_line.strip()
        if not txt:
            continue
        if txt.startswith("# "):
            title = txt.lstrip("# ").strip()
        elif txt.startswith("## "):
            if current:
                sections.append({"heading": current, "content": buf})
            current = txt.lstrip("# ").strip()
            buf = []
        else:
            if current:
                buf.append(txt)
    if current:
        sections.append({"heading": current, "content": buf})
    m = re.search(r"Agent\s+(\d+)", title or path.name)
    agent_id = int(m.group(1)) if m else 0
    return {
        "id": agent_id,
        "title": title or path.stem,
        "sections": sections,
        "source_file": str(path.relative_to(repo_root())),
    }


def main() -> int:
    root = repo_root()
    out_dir = data_dir()
    out_dir.mkdir(parents=True, exist_ok=True)

    # PRD
    prd_candidates = list(root.glob("docs/**/Product Requirements Multi Agent PPM Platform.docx"))
    if not prd_candidates:
        raise FileNotFoundError("Could not find the Product Requirements DOCX under docs/.")
    prd_path = prd_candidates[0]
    requirements = parse_prd(prd_path)
    (out_dir / "requirements.json").write_text(json.dumps(requirements, ensure_ascii=False, indent=2), encoding="utf-8")

    # Agents
    agent_docs = sorted(root.glob("docs_markdown/specs/agents/*.md"))
    agents = [parse_agent_doc(p) for p in agent_docs]
    agents_payload = {
        "source_dir": "docs_markdown/specs/agents/*.md",
        "agents": sorted(agents, key=lambda a: a.get("id", 0)),
    }
    (out_dir / "agents.json").write_text(json.dumps(agents_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Wrote {out_dir / 'requirements.json'}")
    print(f"Wrote {out_dir / 'agents.json'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
